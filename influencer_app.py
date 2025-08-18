from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
import io, os

app = Flask(__name__)

# Hotel configuration mapping
HOTEL_CONFIG = {
    "Sonesta Redondo Beach & Marina": {
        "address": "300 North Harbor Drive Redondo Beach, CA 90277",
        "company_legal_name": "SONESTA REDONDO BEACH LLC"
    },
    "The Allegro Royal Sonesta Hotel": {
        "address": "171 West Randolph Street Chicago, IL 60601",
        "company_legal_name": "SONESTA CHICAGO LLC"
    },
    "The Clift Royal Sonesta Hotel": {
        "address": "495 Geary Street San Francisco, CA 94102",
        "company_legal_name": "Sonesta Clift LLC"
    },
    "Royal Sonesta Chicago Riverfront": {
        "address": "71 East Wacker Drive Chicago, IL 60601",
        "company_legal_name": "Sonesta Chicago LLC"
    },
    "Royal Sonesta Chicago River North": {
        "address": "505 North State Street Chicago, IL 60654",
        "company_legal_name": "Sonesta State Street LLC"
    },
    "Sonesta Hamilton Park Morristown": {
        "address": "175 Park Avenue Florham Park, NJ 07932",
        "company_legal_name": "Sonesta NJ LLC"
    },
    "Sonesta Simply Suites Jersey City": {
        "address": "21 2nd Street Jersey City, NJ 07302",
        "company_legal_name": "Sonesta Jersey City LLC"
    },
    "Sonesta Simply Suites Parsippany Morris Plains": {
        "address": "100 Candlewood Drive Morris Plains, NJ 07950",
        "company_legal_name": "Sonesta Morris Plains LLC"
    },
    "Sonesta ES Suites Parsippany Morris Plains": {
        "address": "3 Gatehall Drive Parsippany-Troy Hills, NJ 07054",
        "company_legal_name": "Sonesta Gatehall Drive LLC"
    },
    "Sonesta Simply Suites Nanuet": {
        "address": "20 Overlook Boulevard Nanuet, NY 10954",
        "company_legal_name": "Sonesta Nanuet LLC"
    },
    "The Yorkville Royal Sonesta Hotel": {
        "address": "220 Bloor Street West Toronto, ON M5S 1T8",
        "company_legal_name": "Sonesta Toronto ULC"
    },
    "Sonesta ES Suites Toronto": {
        "address": "355 South Park Road Thornhill, ON L3T 7W2",
        "company_legal_name": "Sonesta Canada ULC"
    },
    "Royal Sonesta San Juan": {
        "address": "5961 Isla Verde Avenue San Juan, PR 00979",
        "company_legal_name": "Sonesta San Juan LLC"
    }
}


# Map each placeholder in your template to a form field name
PLACEHOLDERS = {
    "[influencer_name]":         "influencer_name",
    "[nights]":                  "nights",
    "[accommodation_dates]":     "accommodation_dates",
    "[room_types]":              "room_types",
    "[check_in_time]":           "check_in_time",
    "[check_in_date]":           "check_in_date",
    "[check_out_time]":          "check_out_time",
    "[check_out_date]":          "check_out_date",
    "[social_media_handles]":    "social_media_handles",
    "[term_dates]":              "term_dates",
    "[hotel_name]":              "hotel_name",
    "[hotel_address]":           "hotel_address",
    "[program_name]":            "program_name",
    "[min_posts_per_day]":       "min_posts_per_day",
    "[pre_post_days]":           "pre_post_days",
    "[features]":                "features",
    "[promotions]":              "promotions",
    "[hotel_handle]":            "hotel_handle",
    "[brand_handle]":            "brand_handle",
    "[company_handle]":          "company_handle",
    "[campaign_hashtag]":        "campaign_hashtag",
    "[NUMBER]":                  "conf_years",
    "[INFLUENCER LEGAL NAME]":   "influencer_name",
    "[HOTEL LEGAL ENTITY NAME]": "company_legal_name",
    "[company_legal_name]":      "company_legal_name",
    "2025-08-20":               "check_in_date",
    "2025-08-25":               "check_out_date",
}

def set_font_formatting(run, font_name="Times New Roman", font_size=11):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def format_paragraph_font(para, font_name="Times New Roman", font_size=11):
    for run in para.runs:
        set_font_formatting(run, font_name, font_size)

@app.route("/", methods=["GET"])
def form():
    return render_template("form.html", hotel_options=list(HOTEL_CONFIG.keys()))

@app.route("/generate", methods=["POST"])
def generate():
    tpl_path = os.path.join(os.path.dirname(__file__), "Influencer_Agreement_Template.docx")
    doc = Document(tpl_path)

    selected_hotel = request.form.get("hotel_name")
    if selected_hotel in HOTEL_CONFIG:
        hotel_address = HOTEL_CONFIG[selected_hotel]["address"]
        company_legal_name = HOTEL_CONFIG[selected_hotel]["company_legal_name"]
    else:
        hotel_address = request.form.get("hotel_address", "")
        company_legal_name = "SONESTA INTERNATIONAL HOTELS CORPORATION"

    # Helper to get ordinal suffix
    def get_day_with_suffix(day):
        return f"{day}{'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')}"

    data = {}
    for ph, field in PLACEHOLDERS.items():
        if field == "hotel_address":
            data[ph] = hotel_address
        elif field == "company_legal_name":
            data[ph] = company_legal_name
        elif field == "influencer_name":
            data[ph] = request.form.get(field, "").upper()
        elif field in ["features", "promotions"]:
            raw = request.form.get(field, "")
            items = [item.strip() for item in raw.split(",") if item.strip()]
            data[ph] = "\n".join(items)
        elif field == "check_in_date":
            check_in_date = request.form.get("check_in_date", "")
            try:
                in_dt = datetime.strptime(check_in_date, "%Y-%m-%d")
                # Format as "August 20, 2025"
                formatted_date = in_dt.strftime("%B %d, %Y").replace(" 0", " ")
                data[ph] = formatted_date
            except Exception:
                data[ph] = check_in_date
        elif field == "check_out_date":
            check_out_date = request.form.get("check_out_date", "")
            try:
                out_dt = datetime.strptime(check_out_date, "%Y-%m-%d")
                # Format as "August 25, 2025"
                formatted_date = out_dt.strftime("%B %d, %Y").replace(" 0", " ")
                data[ph] = formatted_date
            except Exception:
                data[ph] = check_out_date
        elif field == "accommodation_dates":
            check_in_date = request.form.get("check_in_date", "")
            check_out_date = request.form.get("check_out_date", "")
            try:
                in_dt = datetime.strptime(check_in_date, "%Y-%m-%d")
                out_dt = datetime.strptime(check_out_date, "%Y-%m-%d")

                # Format the date range as "August 20–25th, 2025"
                month = in_dt.strftime('%B')
                start_day = in_dt.day
                end_day_with_suffix = get_day_with_suffix(out_dt.day)
                year = out_dt.year

                # This is the bolded date range that will replace the underlined text
                data[ph] = f"**{month} {start_day}–{end_day_with_suffix}, {year}**"
            except Exception:
                data[ph] = ""
        else:
            data[ph] = request.form.get(field, "")

    def replace_placeholders_in_paragraph(para):
        text = para.text
        
        # Handle the specific case of replacing underlined text and removing "To be used"
        if "Check-in at 3:00PM on August 20, 2025" in text:
            # Replace the entire underlined section with just the bolded date range
            text = text.replace("Check-in at 3:00PM on August 20, 2025", data.get("[accommodation_dates]", ""))
            # Remove "To be used" text
            text = text.replace("To be used ", "")
            
            # Clear existing runs and create new ones
            for run in para.runs:
                run._element.getparent().remove(run._element)
            
            # Add the new text
            run = para.add_run(text)
            set_font_formatting(run)
            
        elif any(ph in text for ph in PLACEHOLDERS):
            for run in para.runs:
                run._element.getparent().remove(run._element)

            for ph in PLACEHOLDERS:
                text = text.replace(ph, data[ph])

            lines = text.splitlines()
            for i, line in enumerate(lines):
                run = para.add_run(line)
                set_font_formatting(run)
                if i < len(lines) - 1:
                    run.add_break()
        else:
            format_paragraph_font(para)

    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para)

    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for para in hdr.paragraphs:
                replace_placeholders_in_paragraph(para)
            for table in hdr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_placeholders_in_paragraph(para)

    def ensure_font_formatting(element):
        for para in element.paragraphs:
            format_paragraph_font(para)
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    ensure_font_formatting(cell)

    ensure_font_formatting(doc)
    for section in doc.sections:
        ensure_font_formatting(section.header)
        ensure_font_formatting(section.footer)

    virtual_doc = io.BytesIO()
    doc.save(virtual_doc)
    virtual_doc.seek(0)
    return send_file(
        virtual_doc,
        as_attachment=True,
        download_name="Influencer_Agreement_Filled.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    app.run(debug=True)
